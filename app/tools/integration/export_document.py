"""
Export Document Tool — Export content as formatted PDF, DOCX, Markdown, or HTML.
Uses markdown→HTML conversion + PyMuPDF for PDF generation.
"""

import json
import logging
import base64
import re

from langchain_core.tools import tool
from app.tools.base import nurav_tool, ToolMetadata, ToolStatus, ToolExample

logger = logging.getLogger(__name__)


def _markdown_to_html(md: str, title: str = "", template: str = "report") -> str:
    """Convert markdown to styled HTML."""
    # Headers
    html = md
    for i in range(6, 0, -1):
        html = re.sub(rf'^{"#" * i}\s+(.+)$', rf'<h{i}>\1</h{i}>', html, flags=re.MULTILINE)
    html = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', html)
    html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
    html = re.sub(r'```(\w*)\n(.*?)```', r'<pre><code class="\1">\2</code></pre>', html, flags=re.DOTALL)
    html = re.sub(r'`(.+?)`', r'<code>\1</code>', html)
    html = re.sub(r'\[(.+?)\]\((.+?)\)', r'<a href="\2">\1</a>', html)
    html = re.sub(r'^- (.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'^(\d+)\. (.+)$', r'<li>\2</li>', html, flags=re.MULTILINE)
    # Blockquotes
    html = re.sub(r'^> (.+)$', r'<blockquote>\1</blockquote>', html, flags=re.MULTILINE)
    # Horizontal rules
    html = re.sub(r'^---+$', r'<hr>', html, flags=re.MULTILINE)
    # Paragraphs
    paragraphs = html.split("\n\n")
    result_parts = []
    for p in paragraphs:
        p = p.strip()
        if p and not any(p.startswith(tag) for tag in ("<h", "<pre", "<li", "<blockquote", "<hr")):
            p = f"<p>{p}</p>"
        result_parts.append(p)
    body = "\n".join(result_parts)

    styles = {
        "report": """
            body { font-family: 'Georgia', serif; max-width: 800px; margin: 40px auto; padding: 0 40px; color: #1a1a1a; line-height: 1.7; }
            h1 { font-size: 2em; border-bottom: 3px solid #4F46E5; padding-bottom: 10px; color: #1e1b4b; }
            h2 { font-size: 1.5em; color: #312e81; border-bottom: 1px solid #e5e7eb; padding-bottom: 5px; }
            h3 { color: #4338ca; }
            code { background: #f3f4f6; padding: 2px 6px; border-radius: 4px; font-family: monospace; }
            pre { background: #1e293b; color: #e2e8f0; padding: 16px; border-radius: 8px; overflow-x: auto; }
            blockquote { border-left: 4px solid #4F46E5; margin: 0; padding: 8px 16px; background: #f8f7ff; }
            a { color: #4F46E5; }
            hr { border: none; border-top: 1px solid #e5e7eb; }
        """,
        "paper": """
            body { font-family: 'Times New Roman', serif; max-width: 700px; margin: 60px auto; padding: 0 60px; font-size: 12pt; line-height: 1.5; }
            h1 { text-align: center; font-size: 18pt; }
            h2 { font-size: 14pt; }
            h3 { font-size: 12pt; }
        """,
        "notes": """
            body { font-family: -apple-system, sans-serif; max-width: 700px; margin: 20px auto; padding: 0 20px; }
            h1 { color: #111; } h2 { color: #333; }
        """,
        "minimal": """
            body { font-family: sans-serif; max-width: 680px; margin: 40px auto; padding: 0 20px; }
        """,
    }

    style = styles.get(template, styles["report"])
    title_html = f"<h1>{title}</h1>\n" if title else ""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title or 'Document'}</title>
<style>{style}</style>
</head>
<body>
{title_html}{body}
</body>
</html>"""


def _html_to_pdf_bytes(html: str) -> bytes:
    """Convert HTML to PDF using PyMuPDF (Story API)."""
    try:
        import fitz
        # Use PyMuPDF's Story API for HTML→PDF
        story = fitz.Story(html=html)
        writer = fitz.DocumentWriter(io_device := fitz.open())
        mediebox = fitz.paper_rect("a4")
        while True:
            device = writer.begin_page(mediebox)
            more, _ = story.place(mediebox + (36, 36, -36, -36))
            story.draw(device)
            writer.end_page()
            if not more:
                break
        writer.close()
        return io_device.write()
    except Exception:
        # Fallback: create simple PDF with text only
        import fitz
        doc = fitz.open()
        page = doc.new_page()
        import re
        text = re.sub(r'<[^>]+>', ' ', html)
        text = re.sub(r'\s+', ' ', text).strip()
        page.insert_text((50, 50), text[:5000], fontsize=11)
        return doc.write()


def _md_to_docx_bytes(md: str, title: str = "") -> bytes:
    """Convert markdown to DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io

        doc = Document()

        if title:
            h = doc.add_heading(title, level=0)
            h.runs[0].font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)

        for line in md.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph()
                continue
            # Headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = min(len(header_match.group(1)), 6)
                doc.add_heading(header_match.group(2), level=level)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r'^\d+\. ', line):
                doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style="List Number")
            else:
                p = doc.add_paragraph()
                # Basic bold/italic
                parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line)
                for part in parts:
                    if part.startswith("**") and part.endswith("**"):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith("*") and part.endswith("*"):
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        p.add_run(re.sub(r'`(.+?)`', r'\1', part))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # python-docx not installed — return placeholder
        raise ImportError("python-docx not installed. Run: pip install python-docx")


@nurav_tool(metadata=ToolMetadata(
    name="export_document",
    description="Export AI responses, research results, or any markdown content as formatted PDF, DOCX, Markdown, or HTML. Multiple templates available: report, paper, notes, minimal.",
    niche="integration",
    status=ToolStatus.ACTIVE,
    icon="download",
    version="1.0.0",
    examples=[
        ToolExample(
            input={"content": "# Research Summary\n\nKey findings...", "format": "html", "title": "Research Report", "template": "report"},
            output='{"filename": "research-report.html", "format": "html", "size_kb": 5.2, "file_base64": "..."}',
            description="Export markdown as styled HTML",
        ),
    ],
    input_schema={"content": "str (markdown)", "format": "str ('pdf'|'docx'|'md'|'html')", "title": "str (optional)", "template": "str ('report'|'paper'|'notes'|'minimal')"},
    output_schema={"file_base64": "str (base64)", "filename": "str", "format": "str", "size_kb": "float"},
    avg_response_ms=3000,
    success_rate=0.90,
))
@tool
async def export_document(content: str, format: str = "html", title: str = "", template: str = "report") -> str:
    """Export content as a formatted document."""
    if not content.strip():
        return json.dumps({"error": "No content provided."})

    fmt = format.lower().strip()
    safe_title = re.sub(r'[^\w\s-]', '', title or "document").strip().replace(" ", "-").lower() or "document"

    try:
        if fmt == "md":
            file_bytes = content.encode("utf-8")
            filename = f"{safe_title}.md"

        elif fmt == "html":
            html = _markdown_to_html(content, title, template)
            file_bytes = html.encode("utf-8")
            filename = f"{safe_title}.html"

        elif fmt == "pdf":
            html = _markdown_to_html(content, title, template)
            file_bytes = _html_to_pdf_bytes(html)
            filename = f"{safe_title}.pdf"

        elif fmt == "docx":
            file_bytes = _md_to_docx_bytes(content, title)
            filename = f"{safe_title}.docx"

        else:
            return json.dumps({"error": f"Unknown format: '{fmt}'. Use: pdf, docx, md, html"})

        encoded = base64.b64encode(file_bytes).decode("utf-8")
        return json.dumps({
            "file_base64": encoded,
            "filename": filename,
            "format": fmt,
            "size_kb": round(len(file_bytes) / 1024, 2),
            "template": template,
        })

    except ImportError as e:
        return json.dumps({"error": str(e), "suggestion": "Try format='html' which has no extra dependencies."})
    except Exception as e:
        return json.dumps({"error": f"Export failed: {str(e)}"})
